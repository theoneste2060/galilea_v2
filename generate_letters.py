from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_letterhead(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("[ INSERT NEXVIREON LTD LOGO HERE — approx. 1.5×1.5 inches ]")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.italic = True

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("NexVireon LTD  |  www.nexvireon.com  |  Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 120, 120)

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

def sp(doc, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, after=6, before=0, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = Pt(15)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"

def heading(doc, text, level=1, sb=14, sa=6):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(sb)
    h.paragraph_format.space_after = Pt(sa)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(6, 21, 40)
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(6, 21, 40)

def create(recipient, title, username, sections, filename):
    doc = Document()
    set_margins(doc)
    add_letterhead(doc)
    add_footer(doc)

    sp(doc, "[ DATE — e.g., 17 June 2026 ]", size=10, after=20)
    sp(doc, f"Ref: GGL-REV-{username.upper()}-2026", size=10, after=8, color=RGBColor(120,120,120))

    sp(doc, recipient, bold=True, size=11, after=1)
    sp(doc, title, size=11, after=1)
    sp(doc, "NexVireon LTD", size=11, after=1)
    sp(doc, "[ COMPANY ADDRESS ]", size=11, after=1)
    sp(doc, "[ CITY, COUNTRY ]", size=11, after=20)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("SUBJECT: Formal Request for Website Review — Galilea Global Logistics (www.galileagloballogistics.rw)")
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    sp(doc, f"Dear {recipient.split()[0]},", size=11, after=10)
    sp(doc, "I hope this letter finds you well.", size=11, after=8)

    sp(doc,
        "I am writing on behalf of the NexVireon LTD technology team to formally request your "
        "review and evaluation of the Galilea Global Logistics website, which has been developed "
        "and deployed at https://www.galileagloballogistics.rw/. As Chief Technology Officer, I "
        "oversee the technical delivery of this project, and your feedback as a key stakeholder "
        "is essential to ensure the platform meets both operational requirements and the highest "
        "standards of quality before full handover.",
        size=11, after=10)

    sp(doc,
        "The website is a full-featured logistics platform designed to serve Galilea Global "
        "Logistics' clients across East Africa, China, and international markets. It encompasses "
        "a public-facing corporate website with service catalogues, real-time shipment tracking, "
        "quote management, and content management capabilities through a secure administrative "
        "backend. Your login credentials have been provisioned as follows:",
        size=11, after=10)

    sp(doc, f"     Username: {username}", size=11, after=2)
    sp(doc, "     Password: Galilea@2025", size=11, after=2)
    sp(doc, "     Admin URL: https://www.galileagloballogistics.rw/admin.php", size=11, after=14)

    sp(doc,
        "I kindly request that you review each of the following functional areas in detail. "
        "For each item, please verify that the feature operates correctly, the content is "
        "accurate and appropriately presented, and the user experience meets expectations.",
        size=11, after=12)

    heading(doc, "Detailed Functional Review Checklist", level=1)

    for section_heading, items in sections:
        heading(doc, section_heading, level=2)
        for item in items:
            bullet(doc, item)

    heading(doc, "Reporting Your Findings", level=1)
    sp(doc,
        "Please document any issues, suggestions, or observations in a brief report and "
        "return it to me via email at [YOUR EMAIL ADDRESS] or share it directly with the "
        "development team. For each issue identified, kindly include:",
        size=11, after=8)
    bullet(doc, "The page URL or section where the issue occurs")
    bullet(doc, "A clear description of the problem or suggested improvement")
    bullet(doc, "The priority level (Critical / Major / Minor / Enhancement)")
    bullet(doc, "A screenshot or screen recording if applicable")

    sp(doc, "", after=8)
    sp(doc,
        "Your input will directly inform our final refinement cycle before the platform is "
        "formally handed over to the Galilea Global Logistics operations team. We value your "
        "expertise and look forward to your feedback.",
        size=11, after=12)

    heading(doc, "Timeline", level=1)
    sp(doc,
        "We would appreciate receiving your feedback within [NUMBER] business days from the "
        "date of this letter. Should you require any clarification or technical assistance "
        "during the review process, please do not hesitate to contact me directly.",
        size=11, after=14)

    sp(doc, "Thank you for your time and continued commitment to excellence.", size=11, after=20)
    sp(doc, "Yours sincerely,", size=11, after=30)
    sp(doc, "Nzakizwanimana Theoneste", bold=True, size=12, after=2)
    sp(doc, "Chief Technology Officer", size=11, after=1)
    sp(doc, "NexVireon LTD", size=11, after=1)
    sp(doc, "Email: [YOUR EMAIL ADDRESS]", size=10, after=1, color=RGBColor(80,80,80))
    sp(doc, "Phone: [YOUR PHONE NUMBER]", size=10, after=1, color=RGBColor(80,80,80))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Enclosure: Website login credentials (as above)")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(120,120,120)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Note: Please insert the NexVireon LTD company logo in the header area marked above. "
        "The recommended logo size is approximately 1.5 x 1.5 inches, positioned on the left side "
        "of the letterhead, aligned with the sender information."
    )
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(150,150,150)

    doc.save(filename)

# --- Cedric ---
cedric_sections = [
    ("1. Public-Facing Website — Core Pages", [
        "Homepage: hero section (headline, subtitle, CTA buttons), service highlights grid, 'Why Galilea' value proposition cards, shipment tracking widget, client testimonials carousel, newsletter email sign-up form, footer with all links",
        "Services List page: grid layout of all 8 services with images, titles, short descriptions, and 'Learn More' links",
        "Service Detail pages (x8): full service description (3-5 paragraphs), image, sidebar with benefits checklist and contact phone number, 'Related Services' section, dual CTA buttons (Get a Quote / Contact Us)",
        "Track & Trace page: tracking input field, search button, three info cards (Real-Time Visibility, Milestone History, Proactive Alerts), 'What to Track' identifier badges section, results display area",
        "Contact page: contact form with validation, office information (Rwanda and China addresses, phone, email), 'Why Submit an Inquiry' checklist, sidebar with office cards",
        "Insights / News list page: article cards with images, categories, publication dates, pagination",
        "Insight Detail page: full article layout with cover image, publication metadata, social sharing, 'More Insights' related section",
        "CMS Pages: About Us, Careers, Privacy Policy, Terms of Service, Cookie Policy - verify content accuracy, formatting, and links",
        "404 error page: user-friendly design with navigation back to homepage",
    ]),
    ("2. Navigation & Search", [
        "Desktop top navigation: Home, Services (dropdown), Company (mega-menu with Solutions / Industries / Company columns), Track & Trace - all items centered",
        "Mega-menu dropdown: full-width panel, three-column grid layout (Solutions, Industries, Company), each column with heading, links, and subtitles; Company column includes a promo card",
        "Mobile navigation: hamburger menu toggle, slide-out panel with all nav links, accordion dropdowns for nested items, Sign In button",
        "Site-wide search functionality: search bar visible on all pages, results page with highlighted matches, no-index for search results",
        "Breadcrumb navigation on interior pages",
        "Footer links: all columns (Services, Company, Legal) link to correct pages; newsletter subscription form functions",
    ]),
    ("3. Forms & User Interaction", [
        "Contact / Quote form: all fields (name, email, phone, company, service interest dropdown, message), client-side and server-side validation, successful submission shows confirmation message and stores in database",
        "Newsletter sign-up: email input with validation, successful submission stores subscriber in database, duplicate email handling",
        "Shipment tracking: accepts container number, booking reference, or bill of lading; displays tracking milestones and current status",
    ]),
    ("4. Responsive Design & Cross-Browser Compatibility", [
        "Desktop (1024px+): full layout with mega-menus, multi-column grids, all hover effects",
        "Tablet (768-1023px): collapsed navigation, adjusted grids, touch-friendly tap targets",
        "Mobile (<768px): hamburger menu, stacked single-column layouts, readable font sizes, adequate touch spacing",
        "Cross-browser check: latest Chrome, Firefox, Safari, and Edge - no layout breakage or JavaScript errors",
    ]),
    ("5. Administrative Backend (Admin Panel)", [
        "Login at /admin.php: login form with username/password fields, error messages for invalid credentials, two-factor authentication (TOTP) flow if enabled, 'Back to website' link",
        "Dashboard: summary cards (total inquiries, subscribers, services, pages), recent inquiries list, quick links to manage content",
        "Services CRUD: list, create, edit, delete, reorder services; all fields save and display correctly",
        "Pages CMS: edit CMS pages (About, Careers, Privacy, Terms, Cookies) with rich text editor; content renders correctly on front-end",
        "Menus: view and manage menu structure; parent/child relationships, column group assignment for mega-menu",
        "Settings: general settings (site title, description, logo), contact information (Rwanda/China addresses, phones, emails, social links), email/SMTP configuration, appearance settings, homepage hero slides management",
        "Inquiries: view, filter by status/date, search, export to CSV, mark as read/new, delete",
        "Newsletter Subscribers: view, filter by source, search, export to CSV, delete",
        "Media Library: upload images, view gallery, delete images",
        "User Management (Superadmin): create, edit, delete admin users; assign roles and section permissions",
        "Activity Log: audit trail of all admin actions with timestamps and user information",
        "Account Settings: change password, enable/disable two-factor authentication (TOTP)",
        "Backup & Export: download database backup, export inquiries and subscribers as CSV",
    ]),
    ("6. Performance & Security", [
        "Page load speed: all pages render within acceptable time; images are optimised",
        "HTTPS: site loads securely, no mixed-content warnings",
        "Security headers: X-Content-Type-Options, X-Frame-Options, Referrer-Policy present on all responses",
        "Form validation: both client-side and server-side validation active; SQL injection and XSS protections in place",
        "Session management: admin session timeout, CSRF tokens on all forms",
        "Rate limiting: login brute-force protection (account locks after max attempts)",
    ]),
]

create("Shumbusho Cedric", "Chief Operating Officer", "cedric", cedric_sections,
       "/home/theo/Desktop/Final-Versions/galilea_v2/letter_to_cedric.docx")

# --- Jules ---
jules_sections = [
    ("1. Strategic & Brand Review", [
        "Homepage: overall brand messaging, hero section impact, value proposition clarity, 'Why Galilea' differentiators, visual quality and consistency",
        "Service pages: accuracy and completeness of service descriptions, alignment with Galilea's actual service offerings, professional tone and presentation",
        "Company representation: About Us page content, team credibility, office locations (Rwanda, China), company story and mission",
        "Brand consistency: colour scheme, typography, imagery style, tone of voice across all pages - does it reflect the Galilea Global Logistics brand identity?",
    ]),
    ("2. Customer-Facing Functionality", [
        "Shipment tracking tool: ease of use, clarity of results, usefulness of information displayed",
        "Contact and inquiry process: is it straightforward for a potential client to reach out? Is the form intuitive?",
        "Service discovery: can a visitor easily understand what services Galilea offers and navigate to detailed information?",
        "Call-to-action placement: are CTAs (Get a Quote, Contact Us, Track Shipment) prominent and logically placed throughout the buyer journey?",
    ]),
    ("3. Overall User Experience & Design", [
        "Visual appeal: first impression, design quality, imagery, layout balance",
        "Navigation intuitiveness: can users find what they need within 2-3 clicks?",
        "Mobile experience: site performs well on smartphones and tablets",
        "Content readability: font sizes, line spacing, contrast, paragraph structure",
        "Trust signals: testimonials, partner logos, certifications, office addresses, professional appearance",
    ]),
    ("4. Website Content - CMS Pages", [
        "About Galilea: company history, mission, values, team - is it compelling and accurate?",
        "Careers page: job appeal, company culture messaging",
        "Privacy Policy, Terms of Service, Cookie Policy: accuracy, completeness, legal compliance",
    ]),
    ("5. Administrative Backend (Executive Overview)", [
        "Login and security: two-factor authentication option, secure access",
        "Dashboard: at-a-glance overview of site performance (inquiries, subscribers, content counts)",
        "Settings: ability to update site-wide content (title, description, contact info, social links, hero slides)",
    ]),
]

create("Bawe Shema Jules", "Chief Executive Officer", "Jules", jules_sections,
       "/home/theo/Desktop/Final-Versions/galilea_v2/letter_to_jules.docx")

print("Both letters generated successfully.")
